from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    print("Word library available.")

    DOCX_AVAILABLE = True

except ImportError:
    print("Word library NOT available. Please run: pip install python-docx")

    DOCX_AVAILABLE = False


# 配置

OUTPUT_DIR = "./output"  # 默认在当前目录下创建 output 文件夹

DOCX_NAME = "异构机器人融合平台_v5.1_完整堡垒版.docx"

MD_NAME = "异构机器人融合平台_v5.1_完整堡垒版.md"


# 核心数据

MANTRA = "车道即轨道，信号灯即交警；影子状态机是底线，超时熔断是保险；主备切换不停机，本地缓存不断网；平台管效率，Safety Laser管安全；零信任防作弊，熵增冗余防最坏，因果存证保命。"


THREE_LAWS = [
    (
        "🔒 铁律一：治理层——零信任博弈",
        "任何跨边界的数据默认脏，任何参与者默认是机会主义者。",
        "交叉验证状态（电量 vs 功耗）、信誉度加权调度、防御性编程（假设下游崩溃）。",
    ),
    (
        "🛡️ 铁律二：物理层——熵增冗余",
        "逻辑追求最优，物理必须为最坏买单。",
        "快慢车道物理隔离、时空硬边界（资源生存锁。",
    ),
    (
        "⚖️ 铁律三：生存层——因果存证",
        "事故后 5 分钟内无法还原因果链，系统在法律上非法。",
        "WORM 黑匣子（不可篡改）、全链路因果图谱、降级生存权（断网本地安全停车）。",
    ),
]


RHINOS = [
    (1, "逻辑", "任务分配缺位", "贪心起步 + 信誉度加权"),
    (2, "逻辑", "僵尸占位", "30s 硬超时清理 + 广播释放"),
    (3, "逻辑", "冷启动流量冲击", "启动错峰，注册间隔 ≥ 5s"),
    (4, "逻辑", "级联延迟放大", "滚动时域，只排未来 5s"),
    (5, "逻辑", "并发写冲突", "etcd 分布式锁 / Raft 协议"),
    (6, "逻辑", "饥饿死锁", "任务队列加权轮询 + 老化机制"),
    (7, "逻辑", "版本泥潭", "Adapter SDK 独立版本，承诺 N-1 兼容"),
    (8, "运维", "人工介入风险", "Zone Lockdown + God Mode (Wifi Direct)"),
    (9, "运维", "Adapter 僵尸", "心跳超时 3 次，自动重启进程"),
    (10, "运维", "WORM 衰减", "磁盘容量监控 < 20% 告警，自动转储冷存储"),
    (11, "运维", "知识断层", "etcd 集中配置 + 轮值主备工程师制度"),
    (12, "运维", "验收陷阱", "Ground Truth 校验清单 + 签字确认"),
    (13, "物理", "坐标漂移", "动态安全公式 S = V×K + RTT×V + C_static"),
    (14, "物理", "金属反射", "采购强制选型：带反射强度的激光雷达"),
    (15, "物理", "光照退化", "采购强制选型：工业级相机 + 主动补光"),
    (16, "物理", "地面摩擦系数", "验收强制标准：干燥工况摩擦系数 > 0.4，禁止油污"),
    (17, "安全", "DDS 加密", "默认启用 DDS Security Plugin，性能损耗计入 RTT 冗余"),
    (18, "安全", "身份漂移", "基于 boot_id 监控，重启期间冻结时间窗（乐观等待+悲观超时）"),
]


SPRINTS = [
    ("0", "1-2周", "基础设施：PTP/NTP/DDS-Security/etcd/WORM 部署"),
    ("1", "2-3周", "仿真压测：信号灯/死锁/主备/告警抑制"),
    ("2", "2-3周", "单品牌闭环：影子状态机/熔断/硬编码后退"),
    ("2.5", "2周", "业务网关：WMS/MES 对接（独立部署）"),
    ("3", "3-4周", "多品牌接入：能力向量/车道/身份漂移"),
    ("4", "3-4周", "真机联调：故障注入/VIL/Playback"),
    ("5", "3-4周", "治理层：信誉度/经济模型接口/认证"),
    ("6", "2-3周", "产品化：模板/孪生/Escrow 协议/隐私合规"),
]


TECH_SPECS = [
    (
        "动态安全公式",
        "S = V × K_brake + RTT × V + C_static。约束：T_min 绝对不低于 1.5s（硬下限）。",
    ),
    ("DDS 加密损耗", "加密增加约 15% RTT，已在公式中预留容限。"),
    ("时间同步", "默认 NTP (<10ms)。动态避障场景必须升级 PTP (<1ms)。"),
    (
        "经济模型 (RaaS)",
        "Task Allocator 效用函数预留 gamma 权重（成本因子）。当前 gamma=0.0，未来配置启用。",
    ),
    (
        "隐私合规",
        "数据最小化。视觉数据本地去特征化处理，仅上传骨骼点；日志中人员 ID 使用 Hash 脱敏。",
    ),
]


MONITORING = [
    "全局健康状态（在线/离线/E-Stop）",
    "吞吐量 + 队列深度",
    "E2E 延迟 (RTT)",
    "死锁/冲突计数器",
    "WORM 状态（写入率/容量）",
]


GROUND_TRUTH = [
    "地图匹配度检查（点云 vs 地图，差异 >2㎡ 告警）",
    "地面摩擦系数验收（测试车道满载急停）",
    "地标物理宽度校验",
]


HARDWARE = [
    ("激光雷达", "必须有反射强度输出（防金属反射）"),
    ("相机", "工业级 + 偏振片（防光照退化）"),
    ("地面要求", "摩擦系数 > 0.4，无油污/接缝 >5mm（防打滑）"),
    ("WiFi", "定向天线，通道末端信号 > -65dBm"),
]


# 辅助函数


def set_font(run, name="微软雅黑"):

    try:
        run.font.name = name

        r = run._element

        r.rPr.rFonts.set(qn("w:eastAsia"), name)

    except:
        pass


def set_bg(cell, color):

    try:
        shade = OxmlElement("w:shd")

        shade.set(qn("w:fill"), color)

        cell._element.get_or_add_tcPr().append(shade)

    except:
        pass


def create_word():

    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    try:
        doc.styles["Normal"].font.name = "微软雅黑"

        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    except:
        pass

    # 封面

    t = doc.add_heading("异构机器人融合平台 v5.1 (完整堡垒版)", 0)

    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_font(t.runs[0])

    t.runs[0].font.size = Pt(28)

    t.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    s = doc.add_paragraph("设计、运维与合规交付手册")

    s.alignment = WD_ALIGN_PARAGRAPH.CENTER

    s.runs[0].bold = True

    set_font(s.runs[0])

    d = doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")

    d.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_font(d.runs[0])

    doc.add_page_break()

    # 哲学

    doc.add_heading("核心哲学", 1)

    p = doc.add_paragraph()

    p.add_run('"降维打击"').bold = True

    p.add_run('：通过"固定路径 + 信号灯"简化导航问题。\n')

    p.add_run('"时空分离"').bold = True

    p.add_run("：平台调度与机器人感知（SCS）解耦。")

    set_font(p)

    # 三大铁律

    doc.add_heading("1. 三大铁律（治理层核心）", 1)

    for name, ph, ex in THREE_LAWS:
        p = doc.add_paragraph()

        p.add_run(f"{name}").bold = True

        p.add_run(f"\n哲学：{ph}\n执行：{ex}")

        set_font(p)

    # 灰犀牛表格

    doc.add_heading("2. 18个灰犀牛防御矩阵", 1)

    tbl = doc.add_table(rows=19, cols=4)

    tbl.style = "Light Grid Accent 1"

    h = tbl.rows[0].cells

    hs = ["#", "类别", "陷阱名称", "一句话补丁/策略"]

    for i, x in enumerate(hs):
        h[i].text = x

        h[i].paragraphs[0].runs[0].bold = True

        set_bg(h[i], "4472C4")

        h[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        set_font(h[i].paragraphs[0].runs[0])

    for i, (idx, cat, name, patch) in enumerate(RHINOS):
        row = tbl.rows[i + 1].cells

        row[0].text = str(idx)

        row[1].text = cat

        row[2].text = name

        row[3].text = patch

        bg = "FFFFFF"

        if cat == "安全":
            bg = "FFC7CE"

        elif cat == "物理":
            bg = "FFD966"

        elif cat == "运维":
            bg = "C5E0B4"

        for j, c in enumerate(row):
            set_font(c.paragraphs[0])

            if j > 0:
                set_bg(c, bg)

    # 技术规格

    doc.add_heading("3. 技术规格与边界", 1)

    for k, v in TECH_SPECS:
        p = doc.add_paragraph(style="List Bullet")

        p.add_run(f"{k}：").bold = True

        p.add_run(v)

        set_font(p)

    # Sprint

    doc.add_heading("4. Sprint 实施路线图", 1)

    t_sprint = doc.add_table(rows=9, cols=3)

    t_sprint.style = "Table Grid"

    h_s = ["Sprint", "周期", "核心目标"]

    for i, x in enumerate(h_s):
        t_sprint.rows[0].cells[i].text = x

        t_sprint.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        set_bg(t_sprint.rows[0].cells[i], "4472C4")

        set_font(t_sprint.rows[0].cells[i].paragraphs[0].runs[0])

    for i, (n, p, g) in enumerate(SPRINTS):
        row = t_sprint.rows[i + 1].cells

        row[0].text = n

        row[1].text = p

        row[2].text = g

        for c in row:
            set_font(c.paragraphs[0])

            set_bg(c, "E7E6E6")

    # Runbook

    doc.add_page_break()

    doc.add_heading("5. 运维与应急响应 (Runbook)", 1)

    doc.add_heading("5.1 核心监控 (Grafana)", 2)

    for m in MONITORING:
        doc.add_paragraph(m, style="List Bullet")

    doc.add_heading("5.2 SOP-A：全场死锁", 2)

    doc.add_paragraph("1. 一键解锁：系统强制最低优先级车后退 3m。", style="List Number")

    doc.add_paragraph(
        "2. Zone Lockdown：若无效，选中区域冻结，发送 Level 0 E-Stop。", style="List Number"
    )

    doc.add_paragraph(
        "3. 上帝手柄：使用 Wifi Direct 直连底盘，绕过网络物理接管。", style="List Number"
    )

    doc.add_heading("5.6 SOP-E：网络分区与脑裂", 2)

    doc.add_paragraph(
        "1. 客户端降级：心跳丢失 >5s，切换离线安全模式（停车）。", style="List Number"
    )

    doc.add_paragraph("2. 脑裂防护：etcd Lease 机制，无 Lease 禁止写指令。", style="List Number")

    doc.add_heading("5.7 SOP-F：数据中毒与恶意注入", 2)

    doc.add_paragraph("1. 识别：交叉验证引擎检测物理不可能状态（瞬移/超速）。", style="List Number")

    doc.add_paragraph("2. 熔断：DDS 踢出违规 Participant，标记为 Malicious。", style="List Number")

    doc.add_paragraph("3. 隔离：远程 E-Stop，通知安保。", style="List Number")

    doc.add_heading("5.8 Ground Truth 校验清单", 2)

    for item in GROUND_TRUTH:
        doc.add_paragraph(f"- {item}", style="List Bullet")

    # 附录

    doc.add_page_break()

    doc.add_heading("6. 战略附录", 1)

    doc.add_heading("附录 B：产品化与生态", 2)

    p_e = doc.add_paragraph()

    p_e.add_run("供应商 Escrow (生前遗嘱)：").bold = True

    p_e.add_run(
        "\n强制第三方 Adapter 源码托管。触发条件：厂商破产/停服超 6 个月，甲方有权解锁源码自行维护。"
    )

    doc.add_heading("附录 D：硬件部署白皮书", 2)

    t_hw = doc.add_table(rows=5, cols=2)

    t_hw.style = "Light List Accent 1"

    for i, (name, spec) in enumerate(HARDWARE):
        t_hw.rows[i].cells[0].text = name

        t_hw.rows[i].cells[1].text = spec

        set_font(t_hw.rows[i].cells[0].paragraphs[0])

        set_font(t_hw.rows[i].cells[1].paragraphs[0])

    doc.add_heading("附录 F：安全合规与法律责任", 2)

    p_p = doc.add_paragraph()

    p_p.add_run("隐私保护：").bold = True

    p_p.add_run(
        "\n1. 禁止上传原始人脸视频。\n2. 人员日志 Hash 脱敏存储。\n3. 仅事故发生时授权解密。"
    )

    doc.add_heading("执行口诀", 1)

    q = doc.add_paragraph()

    q.add_run(f'"{MANTRA}"')

    q.style = "Intense Quote"

    set_font(q.runs[0])

    return doc


# 主程序


def main():

    out = Path(OUTPUT_DIR)

    out.mkdir(exist_ok=True)

    print(f"Output directory: {out.absolute()}")

    # Markdown

    md_path = out / MD_NAME

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# 异构机器人融合平台 v5.1 (完整堡垒版)\n")

        f.write("## 设计、运维与合规交付手册\n\n")

        f.write(f"> **核心哲学**：\n> {MANTRA}\n\n")

        f.write("---\n\n")

        f.write("### 3. 18个灰犀牛防御矩阵\n\n")

        f.write("| # | 类别 | 陷阱名称 | 一句话补丁/策略 |\n")

        f.write("|:---:|:---|:---|:---|\n")

        for r in RHINOS:
            f.write(f"| **{r[0]}** | {r[1]} | {r[2]} | {r[3]} |\n")

    print(f"✅ Markdown: {md_path.name}")

    # Word

    if DOCX_AVAILABLE:
        docx_path = out / DOCX_NAME

        doc = create_word()

        if doc:
            doc.save(docx_path)

            print(f"✅ Word: {docx_path.name}")

    else:
        print("⚠️  Word skipped (library missing)")


if __name__ == "__main__":
    main()
