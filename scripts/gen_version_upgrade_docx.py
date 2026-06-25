"""
Generate: D:\EWM ROBOT\REFERENCE\DESIGN\version-upgrade-verification-v3.4-to-v5.0.docx
Includes: EWM + WM dual-system architecture section
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    hs.font.name = 'Microsoft YaHei'

# ── Helpers ──
def add_shaded_cell(cell, color='D9E2F3'):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def make_table(ws, headers, rows, col_widths=None):
    tbl = ws.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell(cell, h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0xFF, 0xFF, 0xFF))
        add_shaded_cell(cell, '2F5496')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            set_cell(cell, str(val), bold=(ci == 0), size=9)
            if ri % 2 == 1:
                add_shaded_cell(cell, 'F2F2F2')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl

def add_bullet(para, label, text, label_color=None):
    run = para.add_run(label)
    run.bold = True
    if label_color:
        run.font.color.rgb = RGBColor(*label_color)
    para.add_run(text)

# ═══════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run('版本升级验证')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SAP EWM 机器人调度平台  v3.4 → v5.0')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
meta = [
    '版本日期：2026-06-25',
    '当前版本：v3.4（生产就绪 · EWM + WM 双系统对接）',
    '目标版本：v5.0（Kubernetes 容器编排）',
    '依据文档：implementation-roadmap-plan · phase3-completion-report · v3.35.md',
]
for m in meta:
    run = p.add_run(m + '\n')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_page_break()

# ═══════════════════════════════════════
# 1. 版本栈对照
# ═══════════════════════════════════════
doc.add_heading('一、版本栈对照', level=1)

h = ['维度', 'v3.4 (当前·生产就绪)', 'v4.0 目标', 'v4.1 目标', 'v5.0 目标']
rows = [
    ['SAP 对接',   'EWM (OData) + WM (RFC) 双系统', '同 v3.4',         '同 v3.4',            '同 v3.4 + gRPC'],
    ['存储引擎',   'SQLite + WAL',                  'PostgreSQL 15+',  'PostgreSQL 15+',     'PostgreSQL HA'],
    ['编排方式',   'Docker Compose',                'Docker Compose',  'Docker Compose',     'Kubernetes'],
    ['机器人接入', 'Node-RED Sub-flow + 策略模式',  'Node-RED Sub-flow', '策略模式 (TS)',    '策略模式 + gRPC'],
    ['仓库支持',   '单仓库 (EWM/WM 混接)',          '多仓库 Plugin',   '多仓库',             '多集群·多仓库'],
    ['扩展性',     '垂直扩展（单机）',              '垂直扩展',        '垂直扩展',           '水平扩展（HPA）'],
    ['部署方式',   'docker compose up -d',           '同 v3.4',         '同 v3.4',            'Helm Chart + ArgoCD'],
    ['监控',       'Prometheus + Grafana',           '同 v3.4',         '同 v3.4',            '+ Loki + Tempo + OTel'],
    ['更新策略',   '停止→重启（有中断）',           '停止→重启',       '停止→重启',          '滚动更新·零停机'],
    ['配置管理',   '.env + 挂载文件',                '同 v3.4',         '同 v3.4',            'ConfigMap + Secret'],
    ['服务发现',   'Docker DNS',                    '同 v3.4',         '同 v3.4',            'CoreDNS + Service Mesh'],
]
make_table(doc, h, rows, col_widths=[3.5, 4.5, 4, 3.5, 4.5])

# ═══════════════════════════════════════
# 2. 双系统架构
# ═══════════════════════════════════════
doc.add_heading('二、核心能力：SAP EWM + SAP WM 双系统对接', level=1)

p = doc.add_paragraph()
run = p.add_run('v3.4 已实现')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run(' — 同时对接 SAP EWM（Extended Warehouse Management）和 SAP WM（Classic Warehouse Management），在同一平台统一调度机器人。')

doc.add_heading('后端插件注册表', level=2)
doc.add_paragraph('采用与机器人策略模式完全相同的 Plugin Registry 模式，每个后端实现统一的 WarehouseBackend 接口：')

h = ['组件', '路径', '职责']
rows = [
    ['WarehouseBackend (ABC)', 'backends/base.py',       '抽象契约：list_tasks · get_task · create_task · confirm_task · cancel_task'],
    ['EwmBackend',             'backends/ewm_backend.py', 'EWM 通过 OData API 对接，RF 扫码确认，仓位管理'],
    ['WmBackend',              'backends/wm_backend.py',  'WM 通过 RFC/BAPI 对接，传统 Transfer Order 处理'],
    ['BackendRegistry',        'backends/registry.py',    '按 backend_type 注册 + 查找，单例'],
    ['Factory',                'backends/factory.py',     '根据 config.yaml 的 sap.warehouses[].backend 选择后端'],
]
make_table(doc, h, rows, col_widths=[5, 5.5, 9])

doc.add_heading('配仓配置示例', level=2)
p = doc.add_paragraph()
p.style.font.size = Pt(9)
p.add_run('config.yaml 中按仓库配置 backend 类型，同一个 sap-bridge 容器可同时对接 EWM 和 WM：').font.size = Pt(9)

# Config table
h = ['仓库', 'backend', '连接方式', '鉴权']
rows = [
    ['WH01 自动化立库', 'ewm', 'odata_url: https://ewm.example.com/...', '${SAP_EWM_AUTH}'],
    ['WH02 传统平库',   'wm',  'rfc_ashost: 10.0.1.100, rfc_sysid: PRD', '${SAP_WM_AUTH}'],
]
make_table(doc, h, rows, col_widths=[4, 2.5, 6, 4.5])

doc.add_heading('已验证能力', level=2)
h = ['能力', 'EWM', 'WM', '说明']
rows = [
    ['创建仓库任务',     'OData POST',    'BAPI_TO_CREATE',   '统一转为 WarehouseTask 模型'],
    ['获取任务列表',     'OData GET',     'BAPI_TO_GET_LIST', '支持分页 + 状态过滤'],
    ['确认完成',         'OData PATCH',   'BAPI_TO_CONFIRM',  'RF 扫码回传'],
    ['取消任务',         'OData DELETE',  'BAPI_TO_CANCEL',   '异常场景回收'],
    ['IDoc 监听',        'XML 解析',      'XML 解析',         '异步大批量订单入口'],
    ['连接健康检查',     'OData /health', 'RFC ping',         'Docker healthcheck 集成'],
]
make_table(doc, h, rows, col_widths=[4, 4, 4, 6.5])

doc.add_heading('对升级路径的影响', level=2)
doc.add_paragraph('EWM + WM 双系统支持已经在 v3.4 完整实现，不在升级路径中。升级中对 SAP 对接层的影响：')
p = doc.add_paragraph()
add_bullet(p, 'v4.0：', '存储从 SQLite 迁到 PostgreSQL，对接逻辑不变', (0xC0, 0x00, 0x00))
p = doc.add_paragraph()
add_bullet(p, 'v4.1：', '策略模式增强，SAP 对接层不变', (0xC0, 0x00, 0x00))
p = doc.add_paragraph()
add_bullet(p, 'v5.0：', 'K8s 编排，对接层作为无状态 Pod 水平扩展，config.yaml 转为 ConfigMap', (0xC0, 0x00, 0x00))

# ═══════════════════════════════════════
# 3. 六大驱动力
# ═══════════════════════════════════════
doc.add_heading('三、为什么要从 v3.4 升级到 v5.0', level=1)

drivers = [
    ('驱动力 1：存储层瓶颈 — SQLite 无法支撑多仓库并发',
     '现状：SQLite 单写入者架构，并发写锁库；Node-RED 与 SAP Bridge 共用，写入竞争导致 checkpoint 卡顿；WAL 模式缓解有限。',
     '升级路径：v3.4 SQLite → v4.0 PostgreSQL → v5.0 PG HA 主从',
     '收益：写入并发 1→N；Outbox 在 PG 事务中更可靠；支持多仓库独立 schema'),
    ('驱动力 2：编排方式限制 — Docker Compose 不是生产级编排',
     '现状：无滚动更新（更新=全部重启有中断）；无应用级自愈；单副本；无配置中心。',
     '升级路径：v3.4 Docker Compose → v5.0 Kubernetes (Deployment+Service+ConfigMap)',
     '收益：滚动更新零停机；Liveness/Readiness Probe 自愈；HPA 扩缩容；ConfigMap/Secret 集中管理'),
    ('驱动力 3：品牌扩展 — Node-RED Sub-flow 维护成本爆炸',
     '现状：每新品牌需拖 Sub-flow；难做单元测试；通用逻辑（状态机/幂等/重试）在 Sub-flow 间重复。',
     '升级路径：v3.4 Sub-flow → v4.1 Python 策略模式（基础已实现）→ v5.0 + gRPC',
     '收益：策略模式可测试（现有 82% 覆盖）；新品牌只需实现 Strategy 接口；消除重复逻辑'),
    ('驱动力 4：部署频率与回滚能力',
     '现状：部署 = docker compose down && up，≥30s 中断；回滚需重启；无法金丝雀发布。',
     '升级路径：v3.4 粗暴重启 → v5.0 Helm 滚动更新',
     '收益：更新期间旧 Pod 继续服务；Helm 一令回滚；金丝雀发布 10% 流量'),
    ('驱动力 5：运维与可观测性不足',
     '现状：日志靠 Docker logs；告警靠飞书 Webhook 无去重/升级；无分布式追踪。',
     '升级路径：v3.4 PG+Grafana → v5.0 +Loki+Tempo+OpenTelemetry',
     '收益：Loki 日志聚合统一看板；Tempo 端到端追踪（SAP→MQTT→机器人）；告警升级/静默/路由'),
    ('驱动力 6：多仓库/多数据中心部署需求',
     '现状：单仓库架构；新仓库=全新部署一套 Docker Compose；跨仓库需自定义桥接。',
     '升级路径：v3.4 单仓库 → v4.0 多仓库 Plugin → v5.0 多集群+Service Mesh',
     '收益：每个仓库独立扩缩容；Istio 统一流量管理；MQTT+Kafka 事件总线'),
]

for title, status, path, benefit in drivers:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    add_bullet(p, '现状：', status)
    p = doc.add_paragraph()
    add_bullet(p, '升级路径：', path, (0xC0, 0x00, 0x00))
    p = doc.add_paragraph()
    add_bullet(p, '收益：', benefit, (0x00, 0x80, 0x00))

# ═══════════════════════════════════════
# 4. 版本验证矩阵
# ═══════════════════════════════════════
doc.add_heading('四、版本验证矩阵', level=1)

doc.add_heading('v3.4 → v4.0：SQLite → PostgreSQL', level=2)
h = ['#', '验证项', '方法', '通过标准', '优先级']
r = [
    ['1', '数据迁移完整性',  'SELECT COUNT(*) + CHECKSUM 对比新旧库',  '行数一致，checksum 匹配',  'P0'],
    ['2', 'Outbox 事务一致性', 'INSERT order + outbox 在同一 PG 事务',  '两条同时写入或都不写',     'P0'],
    ['3', '连接池泄漏',       '连续 100 次 SAP 调用后查 pg_stat_activity', '活跃连接 ≤ 5',            'P0'],
    ['4', '查询性能',         '执行关键 SQL（订单查询/状态统计）',       '执行时间 < 100ms',          'P1'],
    ['5', '回滚能力',         '执行 --rollback 脚本',                    '数据完整恢复到 SQLite',     'P0'],
]
make_table(doc, h, r, col_widths=[1, 4, 6, 5, 2])

doc.add_heading('v4.0 → v4.1：策略模式', level=2)
h = ['#', '验证项', '方法', '通过标准', '优先级']
r = [
    ['1', '品牌策略注册',   'GET /api/strategies',     '返回所有已注册品牌',           'P0'],
    ['2', '策略路由',       '分别下发 Geek+ / HaiRobotics 订单', '各自调用对应策略 dispatch()', 'P0'],
    ['3', '版本兼容',       '各策略 supported_versions', '所有品牌 ≥ VDA5050 v1.1.0',  'P1'],
    ['4', '降级处理',       '发送未注册品牌订单',       '返回 501，日志 unknown_brand',  'P1'],
]
make_table(doc, h, r, col_widths=[1, 4, 6, 5, 2])

doc.add_heading('v4.1 → v5.0：Kubernetes', level=2)
h = ['#', '验证项', '方法', '通过标准', '优先级']
r = [
    ['1', 'Helm 部署',     'helm install robot-platform ./charts',  '所有 Pod Running',         'P0'],
    ['2', '健康检查',      'kubectl get pods',                      '所有 Pod Ready ≥ 3/3',     'P0'],
    ['3', '滚动更新',      'helm upgrade 更新镜像版本',              '零停机，旧 Pod 逐步替换',  'P0'],
    ['4', '自愈测试',      'kubectl delete pod nodered-xxx',        '新 Pod 自动创建 < 30s',    'P0'],
    ['5', 'HPA 扩缩容',   'k6 压测 1000 单并发',                   '自动扩容到 3+ 副本',       'P1'],
    ['6', '配置挂载',      'ConfigMap 更新后 Pod 读取',              '无需重建 Pod',             'P1'],
]
make_table(doc, h, r, col_widths=[1, 4, 6, 5, 2])

# ═══════════════════════════════════════
# 5. 里程碑
# ═══════════════════════════════════════
doc.add_heading('五、阶段实施里程碑', level=1)
h = ['阶段', '时间窗口', '关键任务', '交付物']
r = [
    ['v4.0 存储升级\nSQLite → PG',
     '2026-07-01 → 2026-07-21\n（3 周）',
     '• PostgreSQL 容器部署\n• 数据迁移脚本执行\n• Outbox 表迁移验证\n• Node-RED PG 连接配置更新',
     '• postgresql 容器运行\n• orders/outbox 表在 PG\n• 迁移验证报告'],
    ['v4.1 策略模式\nSub-flow → Python',
     '2026-08-01 → 2026-08-15\n（2 周）',
     '• 品牌策略注册中心实现\n• 现有品牌策略迁移\n• 单元测试与回归验证',
     '• GET /api/strategies 端点\n• 5 品牌策略类\n• 305 测试保持通过'],
    ['v5.0 K8s 编排\nCompose → K8s',
     '2026-09-01 → 2026-09-28\n（4 周）',
     '• K8s 集群搭建与网络配置\n• Docker → Helm Chart 迁移\n• 滚动更新与 HPA 配置\n• 全链路端到端验证',
     '• Helm chart 包\n• 滚动更新验证报告\n• HPA 压力测试报告\n• 回滚演练记录'],
]
make_table(doc, h, r, col_widths=[3.5, 3.5, 5.5, 5])

# ═══════════════════════════════════════
# 6. 风险矩阵
# ═══════════════════════════════════════
doc.add_heading('六、风险与缓解', level=1)
h = ['风险', '概率', '影响', '风险等级', '缓解措施']
r = [
    ['PG 迁移数据不一致',         'LOW',  'HIGH',    'MED',  '迁移前后 COUNT + CHECKSUM 双重校验'],
    ['策略模式重构引入回归',      'MED',  'MED',     'MED',  '现有 305 测试套件作为回归基线'],
    ['K8s 学习曲线延迟部署',      'MED',  'MED',     'MED',  '先小集群验证，再迁移生产'],
    ['滚动更新期间状态丢失',      'LOW',  'CRITICAL','HIGH', 'Redis 持久化 + 订单状态在 PG 中'],
    ['MQTT 桥接在 K8s 中不可靠', 'LOW',  'HIGH',    'MED',  'StatefulSet + Headless Service'],
    ['网络策略导致通信断',        'MED',  'HIGH',    'HIGH', '预先定义 NetworkPolicy，staging 验证'],
    ['存储卷在 K8s 中丢失',       'LOW',  'CRITICAL','HIGH', 'PVC + 定期备份 + 云盘快照'],
]
make_table(doc, h, r, col_widths=[4.5, 2, 2.5, 2.5, 6])

# ═══════════════════════════════════════
# 7. 兼容性保障
# ═══════════════════════════════════════
doc.add_heading('七、兼容性保障', level=1)
h = ['升级路径', '可逆性', '保障措施']
r = [
    ['v3.4 ↔ v4.0\nSQLite ↔ PostgreSQL', '可逆',
     '迁移脚本支持双向 sync；切换前自动完整备份；保留 SQLite 库 7 天'],
    ['v4.0 ↔ v4.1\nSub-flow ↔ 策略模式', '可逆',
     '策略模式与 Sub-flow 并行运行；通过 feature flag 切换；dummy 模式只记录不执行'],
    ['v4.1 → v5.0\nDocker Compose → K8s', '单向（保留回滚环境）',
     'K8s 迁移后保留原 Compose 环境 2 周；Helm 版本管理快速回滚；数据层（PG/Redis）保持不变'],
]
make_table(doc, h, r, col_widths=[5, 4, 8.5])

# ═══════════════════════════════════════
# 8. 当前状态 + 升级触发条件
# ═══════════════════════════════════════
doc.add_heading('八、当前状态 + 升级触发条件', level=1)

doc.add_heading('当前版本状态（v3.4）— 生产就绪', level=2)
h = ['维度', '状态', '验证结果', '证据']
r = [
    ['单元测试',       'PASS',   '305 passed, 0 failed',        '82% line coverage'],
    ['CI/CD 管线',     'PASS',   '5-stage (lint/ADR/test/build/Trivy)', '.github/workflows/ci.yml'],
    ['安全审计',       'PASS',   '0 HIGH · 0 CRITICAL',          'docs/audit-findings-p1.md'],
    ['48h 清单',       'PASS',   '39/39 结构项通过，4 项现场',   'docs/48h-checklist-v3.4.md'],
    ['服务状态',       'HEALTHY','13 容器全部运行',               'docker compose ps'],
    ['SAP 双系统',    'EWM+WM', 'OData + RFC 双后端已实现',      'sap-bridge/backends/{ewm,wm}_backend.py'],
    ['版本标记',       'v3.4',   '全部 VERSION 文件统一',         '.claude/rules/VERSION'],
]
make_table(doc, h, r, col_widths=[3, 3, 5.5, 6])

doc.add_heading('v4.0 迁移触发器（SQLite -> PostgreSQL）', level=2)
h = ['', '触发条件', '阈值', '说明', '检查命令']
r = [
    ['', '新增第二个仓库',      '>= 2 仓库并发',   '多仓库并发写入 SQLite 瓶颈', 'ls /data/warehouse_*'],
    ['', '单日订单量超标',      '> 10,000 单/日', 'WAL checkpoint 延迟 > 2s',  'SELECT COUNT(*)'],
    ['', '需要 Row-Level Security', '多租户隔离需求', 'SQLite 无行级权限控制',  '安全评估报告'],
    ['', 'SLA 要求 >= 99.9%',   '可用性需求',      'Docker Compose 无法满足',    'SLA 合同条款'],
]
make_table(doc, h, r, col_widths=[1, 4, 4, 5, 3.5])

doc.add_heading('v5.0 迁移触发器（Docker Compose -> Kubernetes）', level=2)
h = ['', '触发条件', '阈值', '说明', '检查命令']
r = [
    ['', '机器人数量',       '> 50 台',       '单 Node-RED 实例瓶颈',       'SELECT COUNT(*) FROM robots'],
    ['', '零停机部署需求',    '业务要求无中断', 'Docker Compose 无法滚动更新', '变更管理策略'],
    ['', '多数据中心部署',    '>= 2 数据中心',  '需多集群管理',                '基础设施规划'],
    ['', '容器数量',         '> 30 容器',      'Docker Compose 管理复杂',    'docker ps | wc -l'],
]
make_table(doc, h, r, col_widths=[1, 4, 4, 5, 3.5])

# ═══════════════════════════════════════
# 9. 总结
# ═══════════════════════════════════════
doc.add_heading('九、转换理由总结 — 一句话版', level=1)
h = ['问题', '核心答案']
r = [
    ['v3.4 有什么问题？',                 '没问题。生产就绪，已实现 EWM + WM 双系统对接。升级是"架构演进"，不是"修复缺陷"'],
    ['v3.4 能接 EWM 和 WM 吗？',          '能。双后端已在 v3.4 实现：EWM 走 OData，WM 走 RFC/BAPI，通过 config.yaml 按仓库配置切换'],
    ['为什么要升级？',                    '为了应对未来增长：多仓库、高并发、零停机、多数据中心'],
    ['什么时候升级？',                    '由业务指标触发（仓库数>1/订单量>1万/SLA>99.9%），不强升'],
    ['v4.0 解决什么？',                  'SQLite→PostgreSQL：解决并发写瓶颈、outbox 事务一致性、多租户隔离'],
    ['v4.1 解决什么？',                  'Sub-flow→策略模式：解决品牌扩展测试难、重复逻辑、维护成本'],
    ['v5.0 解决什么？',                  'Docker Compose→K8s：解决滚动更新、自愈、水平扩展、配置中心'],
    ['升级可逆吗？',                      'v3.4↔v4.0 可逆；v4.0↔v4.1 可逆；v4.1→v5.0 单向但保留回滚环境 2 周'],
    ['当前做什么？',                      '聚焦跑稳态生产，写好运行手册和降级演练，等业务增长触发升级'],
]
make_table(doc, h, r, col_widths=[4.5, 13])

# ═══════════════════════════════════════
# 附录
# ═══════════════════════════════════════
doc.add_heading('附录：关键文件索引', level=1)
h = ['文件', '路径', '说明']
r = [
    ['实施路线图',     '.claude/memory/implementation-roadmap-plan.md',    '版本规划、多仓库架构'],
    ['Phase 3 报告',   'docs/phase3-completion-report.md',                 '当前 v3.4 完成状态'],
    ['架构设计 v3.35', 'REFERENCE/DESIGN/SAP-EWM-...v3.35.md',            '全维防御体系设计'],
    ['架构总览 v3.4',  'REFERENCE/CURSOR 3.4 docs/CURSOR_ARCHITECTURE...', '微服务拆分预备'],
    ['仓库后端抽象',   '.claude/memory/warehouse-backend-abstraction.md',  '双 EWM+WM 后端 Plugin 设计'],
    ['双后端代码',     'sap-bridge/backends/',                              'EWM + WM backend 实现'],
    ['CLAUDE.md',      'CLAUDE.md',                                        '项目主配置 v3.4'],
]
make_table(doc, h, r, col_widths=[4, 8, 5.5])

# ── Save ──
out = r'D:\EWM ROBOT\REFERENCE\DESIGN\version-upgrade-verification-v3.4-to-v5.0.docx'
doc.save(out)
print(f'OK Saved: {out}')
print(f'   Size: {os.path.getsize(out):,} bytes')
